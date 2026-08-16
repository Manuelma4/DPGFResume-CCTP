from __future__ import annotations

import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from fastapi.testclient import TestClient
from openpyxl import load_workbook

from app import auth, config, store
from app import directory as directory_module
from app import llm as llm_module
from app.excel_export import export_analysis, export_analysis_files
from app.extractors import extract_document
from app.main import _apply_perimeter_assist, _apply_unit_assist, app
from app.parser import classify_lot_family, lot_sort_key, parse_document, recompute_stats

# Tests must stay fast, deterministic and free of real network calls even
# when a real LIHA token is configured in .env for the running app. Tests
# that specifically exercise the LIHA-assisted paths mock app.llm directly.
config.USE_LLM = False


def make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("LOT 03 — GROS ŒUVRE", level=1)
    document.add_heading("3.1 INFRASTRUCTURES", level=2)
    document.add_heading("3.1.1 Terrassements", level=3)
    document.add_paragraph(
        "Terrassements en pleine masse, évacuation comprise. Unité de règlement : m³."
    )
    document.add_heading("3.1.2 Béton de propreté", level=3)
    document.add_paragraph("Béton dosé suivant prescriptions, quantité 12,5 m².")
    document.add_heading("3.2 OUVRAGES DIVERS", level=2)
    document.add_heading("3.2.1 Réservations et percements", level=3)
    document.add_paragraph("Ensemble des sujétions, au forfait.")
    document.add_heading("3.2.9 Location mensuelle", level=3)
    document.add_paragraph("Installation maintenue pendant le chantier. Unité : mois.")
    document.save(path)


def sample_analysis(lot: dict) -> dict:
    return {
        "id": "a" * 32,
        "project": {
            "name": "BONDUELLE — Bâtiment Gay Lussac",
            "reference": "25_189",
            "client": "BONDUELLE",
            "phase": "PRO",
            "due_date": "",
        },
        "lots": [lot],
        "stats": recompute_stats([lot]),
    }


class ExtractionTests(unittest.TestCase):
    def test_docx_hierarchy_units_and_traceability(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 03.docx"
            make_docx(path)
            extracted = extract_document(path)
            lot = parse_document(extracted, "src_test")

        self.assertEqual(lot["code"], "03")
        self.assertIn("GROS", lot["title"].upper())
        items = [line for line in lot["lines"] if line["kind"] == "item"]
        self.assertGreaterEqual(len(items), 3)
        by_code = {line["code"]: line for line in items}
        self.assertEqual(by_code["3.1.1"]["unit"], "m³")
        self.assertEqual(by_code["3.1.2"]["quantity"], 12.5)
        self.assertEqual(by_code["3.2.9"]["unit"], "mois")
        self.assertIn("source_excerpt", by_code["3.2.1"])

    def test_docx_word_toc_entries_are_not_extracted_as_headings(self) -> None:
        # Word's automatic table of contents applies built-in "TOC N" styles
        # and bakes the resolved page number into the plain text of each
        # entry (e.g. "4.1.1 Ragréage 11"). Left unfiltered, real CCTP files
        # produce a duplicate, page-number-suffixed phantom line for every
        # real heading, positioned before the actual content.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 04 GROS OEUVRE.docx"
            document = Document()
            document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
            document.styles.add_style("TOC 3", WD_STYLE_TYPE.PARAGRAPH)
            document.add_heading("LOT 04 — GROS OEUVRE", level=1)
            document.add_paragraph("4. Description des ouvrages 11", style="TOC 1")
            document.add_paragraph("4.1.1 Ragréage 11", style="TOC 3")
            document.add_heading("4. Description des ouvrages", level=1)
            document.add_heading("4.1 Travaux préparatoires", level=2)
            document.add_heading("4.1.1 Ragréage", level=3)
            document.add_paragraph("Fourniture et pose d'un ragréage autolissant.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_toc")

        designations = [line["designation"] for line in lot["lines"]]
        self.assertNotIn("Description des ouvrages 11", designations)
        self.assertNotIn("Ragréage 11", designations)
        self.assertIn("Ragréage", designations)
        by_code = {line["code"]: line for line in lot["lines"]}
        self.assertEqual(by_code["4.1.1"]["designation"], "Ragréage")

    def test_docx_word_auto_numbered_headings_get_synthesized_codes(self) -> None:
        # Many real CCTP use Word's native multilevel-list numbering, which
        # is computed at render time and is absent from the paragraph's own
        # text — only the Heading/Titre style survives extraction. The parser
        # must still recover a usable hierarchy from the style levels alone.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 20 CARRELAGE.docx"
            document = Document()
            document.add_heading("LOT 20 — CARRELAGE", level=1)
            document.add_heading("Description des ouvrages", level=1)
            document.add_heading("Travaux préparatoires", level=2)
            document.add_heading("Ragréage", level=3)
            document.add_paragraph("Fourniture et mise en œuvre d'un ragréage.")
            document.add_heading("Carrelage", level=2)
            document.add_heading("Carrelage grès cérame 45 x 45", level=3)
            document.add_paragraph("Fourniture et pose d'un carrelage grès cérame.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_style_numbering")

        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Ragréage"]["code"], "1.1.1")
        self.assertEqual(by_designation["Ragréage"]["kind"], "item")
        self.assertEqual(
            by_designation["Carrelage grès cérame 45 x 45"]["code"], "1.2.1"
        )
        self.assertEqual(lot["perimeter"]["method"], "explicit_anchor")

    def test_only_x_and_x_x_codes_are_titles_even_with_numbered_children(self) -> None:
        # Real DPGF only bold-title "x" chapters and "x.x" sub-chapters (e.g.
        # "3" SPECIFICATIONS TECHNIQUES GENERALES, "3.1" TRAVAUX GENERAUX).
        # A "3.4.2" heading that itemises its own sub-parts (3.4.2.1, 3.4.2.2)
        # is still a priceable row, not a spurious bold header — validated
        # against a real Bonduelle LOT 01 VRD CCTP/DPGF pair.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 01 VRD.docx"
            document = Document()
            document.add_heading("LOT 01 — VRD", level=1)
            document.add_heading("3.4 OUVRAGES D'ASSAINISSEMENT", level=2)
            document.add_heading(
                "3.4.2 Fourniture et pose de canalisation", level=3
            )
            document.add_heading("3.4.2.1 Tranchée pour pose de canalisation", level=4)
            document.add_paragraph("Section courante, quantité en ml.")
            document.add_heading("3.4.2.2 Remblai de tranchée", level=4)
            document.add_paragraph("Remblai compacté, quantité en m³.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_deep_title")

        by_code = {line["code"]: line for line in lot["lines"]}
        self.assertEqual(by_code["3.4"]["kind"], "section")
        self.assertEqual(by_code["3.4.2"]["kind"], "item")
        self.assertEqual(by_code["3.4.2.1"]["kind"], "item")
        self.assertEqual(by_code["3.4.2.2"]["kind"], "item")

    def test_curly_apostrophe_matches_straight_apostrophe_rules(self) -> None:
        # Word CCTP text overwhelmingly types "d’étanchéité" with a curly
        # apostrophe (’) while UNIT_RULES/DECOMPOSITION_RULES are written in
        # this codebase with a straight one ('). A ml rule keyed on
        # "releves? d'etancheite" must still fire against the curly form.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 08 COUVERTURE.docx"
            document = Document()
            document.add_heading("LOT 08 — COUVERTURE ETANCHEITE", level=1)
            document.add_heading("Description des ouvrages", level=1)
            document.add_heading("Relevés d’étanchéité", level=2)
            document.add_paragraph("Relevés d’étanchéité en périphérie de toiture.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_apostrophe_unit")

        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Relevés d’étanchéité"]["unit"], "ml")

    def test_oe_ligature_matches_plain_oe_rules(self) -> None:
        # "GROS ŒUVRE" (correct French typography, œ = U+0153) does not
        # decompose to "oe" under NFKD like an accented letter would — found
        # while adding fondations_gros_oeuvre-specific unit rules: the lot
        # family classifier silently returned "autre" for any real CCTP
        # using the ligature, since LOT_FAMILY_RULES is written with plain
        # "oeuvre". Same class of bug as the curly-apostrophe one above.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 03 GROS OEUVRE.docx"
            document = Document()
            document.add_heading("LOT 03 — GROS ŒUVRE", level=1)
            document.add_heading("Description des ouvrages", level=1)
            document.add_heading("Maçonnerie de remplissage", level=2)
            document.add_paragraph("Fourniture et pose, quantité 40.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_oe_ligature")

        self.assertEqual(lot["code"], "03")
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Maçonnerie de remplissage"]["unit"], "m²")

    def test_curly_apostrophe_prevents_duplicate_decomposition_line(self) -> None:
        # Same bug, but for the decomposition dedup check: the CCTP writes
        # out "Vanne d’isolement DN 15" (curly apostrophe) explicitly as its
        # own line, while the DECOMPOSITION_RULES child constant is written
        # "Vanne d'isolement DN 15" (straight). Before the _normalized() fix
        # this mismatch made the dedup check blind to the duplicate.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 16 PLOMBERIE.docx"
            document = Document()
            document.add_heading("LOT 16 — PLOMBERIE", level=1)
            document.add_heading("Description des ouvrages", level=1)
            document.add_heading("Vanne d’isolement", level=2)
            document.add_paragraph("Vanne à boisseau sphérique.")
            document.add_heading("Vanne d’isolement DN 15", level=2)
            document.add_paragraph("Section courante DN 15, quantité 4 U.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_apostrophe_dedup")

        matches = [
            line for line in lot["lines"] if line["designation"] == "Vanne d’isolement DN 15"
        ]
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0]["origin"], "deterministic-v2")

    def test_pdf_page_is_preserved(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "LOT 08 ELECTRICITE.pdf"
            pdf = fitz.open()
            page = pdf.new_page()
            page.insert_text((50, 70), "LOT 08 - ELECTRICITE", fontsize=16)
            page.insert_text((50, 110), "8.1 COURANTS FORTS", fontsize=14)
            page.insert_text((50, 140), "8.1.1 Tableau divisionnaire", fontsize=12)
            page.insert_text((50, 165), "Fourniture et pose de l'ensemble au forfait.", fontsize=10)
            pdf.save(path)
            pdf.close()
            lot = parse_document(extract_document(path), "src_pdf")

        item = next(line for line in lot["lines"] if line["code"] == "8.1.1")
        self.assertEqual(item["source_page"], 1)
        self.assertEqual(item["unit"], "U")

    def test_pdf_reconstructs_numbered_hierarchy_and_rejects_noise(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "LOT 06 SERRURERIE.pdf"
            pdf = fitz.open()
            cover = pdf.new_page()
            cover.insert_text((50, 70), "BONDUELLE - BATIMENT GAY LUSSAC", fontsize=16)
            cover.insert_text((50, 110), "Architecte", fontsize=12)
            toc = pdf.new_page()
            toc.insert_text((50, 60), "SOMMAIRE", fontsize=16)
            toc.insert_text((50, 100), "3. DESCRIPTION DES OUVRAGES ........ 3", fontsize=11)
            work = pdf.new_page()
            work.insert_text((50, 70), "3.", fontsize=14)
            work.insert_text((105, 70), "DESCRIPTION DES OUVRAGES", fontsize=14)
            work.insert_text((50, 115), "3.1", fontsize=12)
            work.insert_text((105, 115), "ESCALIER EXTERIEUR", fontsize=12)
            work.insert_text((50, 145), "Epaisseur 12 mm et hauteur 4.1 m.", fontsize=10)
            work.insert_text((50, 190), "3.2", fontsize=12)
            work.insert_text((105, 190), "OPTION : GARDE-CORPS", fontsize=12)
            pdf.save(path)
            pdf.close()

            lot = parse_document(extract_document(path), "src_split")

        designations = {line["designation"] for line in lot["lines"]}
        self.assertNotIn("BONDUELLE - BATIMENT GAY LUSSAC", designations)
        self.assertNotIn("SOMMAIRE", designations)
        by_code = {line["code"]: line for line in lot["lines"]}
        self.assertEqual(by_code["3.1"]["designation"], "ESCALIER EXTERIEUR")
        self.assertEqual(by_code["3.1"]["quantity"], None)
        self.assertEqual(by_code["3.1"]["unit"], "U")
        self.assertFalse(by_code["3.2"]["included"])
        self.assertEqual(lot["perimeter"]["method"], "explicit_anchor")

    def test_unit_rules_match_real_dpgf_vocabulary(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 06 SERRURERIE.docx"
            document = Document()
            document.add_heading("LOT 06 — SERRURERIE", level=1)
            document.add_heading("6.1 OUVRAGES", level=2)
            document.add_heading("6.1.1 Garde-corps à barreaudage", level=3)
            document.add_paragraph("Fourniture et pose en acier galvanisé thermolaqué.")
            document.add_heading("6.1.2 Sprinkler", level=3)
            document.add_paragraph("Tête de sprinkler montante, fourniture et pose.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_units")

        by_code = {line["code"]: line for line in lot["lines"]}
        # Mined from real Moduo DPGF: garde-corps is overwhelmingly billed per
        # ml, not per unit as the previous keyword list assumed.
        self.assertEqual(by_code["6.1.1"]["unit"], "ml")
        self.assertEqual(by_code["6.1.2"]["unit"], "U")

    def test_unit_override_beats_generic_group(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 16 PLOMBERIE.docx"
            document = Document()
            document.add_heading("LOT 16 — PLOMBERIE", level=1)
            document.add_heading("16.1 OUVRAGES", level=2)
            document.add_heading("16.1.1 Isolation acoustique", level=3)
            document.add_paragraph("Traitement acoustique forfaitaire du local technique.")
            document.add_heading("16.1.2 Vanne d'isolement", level=3)
            document.add_paragraph("Vanne à boisseau sphérique, fourniture et pose.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_override")

        by_code = {line["code"]: line for line in lot["lines"]}
        # "isolation" alone defaults to m² (surface work), but real DPGF data
        # shows "isolation acoustique" is billed as a lump-sum Ens package.
        self.assertEqual(by_code["16.1.1"]["unit"], "Ens")
        self.assertEqual(by_code["16.1.2"]["unit"], "U")

    def test_unit_depends_on_lot_family_for_ambiguous_words(self) -> None:
        # "câblage" is globally ambiguous (49-73% either way across trades in
        # the real corpus) but pure within a single trade: ~ml in VRD (buried
        # cable trenches, billed by length) vs ~Ens in électricité (a lump
        # cabling package). Same keyword, different lot, different unit.
        with tempfile.TemporaryDirectory() as directory:
            vrd_path = Path(directory) / "CCTP LOT 02 VRD.docx"
            vrd_doc = Document()
            vrd_doc.add_heading("LOT 02 — VRD", level=1)
            vrd_doc.add_heading("2.1 DESCRIPTION DES OUVRAGES", level=2)
            vrd_doc.add_heading("2.1.1 Câblage basse tension", level=3)
            vrd_doc.add_paragraph("Câblage enterré en tranchée commune.")
            vrd_doc.save(vrd_path)
            vrd_lot = parse_document(extract_document(vrd_path), "src_vrd")

            elec_path = Path(directory) / "CCTP LOT 14 ELECTRICITE.docx"
            elec_doc = Document()
            elec_doc.add_heading("LOT 14 — ELECTRICITE CFO/CFA", level=1)
            elec_doc.add_heading("14.1 DESCRIPTION DES OUVRAGES", level=2)
            elec_doc.add_heading("14.1.1 Câblage basse tension", level=3)
            elec_doc.add_paragraph("Câblage des équipements terminaux.")
            elec_doc.save(elec_path)
            elec_lot = parse_document(extract_document(elec_path), "src_elec")

        vrd_item = next(
            line for line in vrd_lot["lines"] if line["designation"] == "Câblage basse tension"
        )
        elec_item = next(
            line for line in elec_lot["lines"] if line["designation"] == "Câblage basse tension"
        )
        self.assertEqual(vrd_item["unit"], "ml")
        self.assertEqual(elec_item["unit"], "Ens")

    def test_generic_administrative_headings_are_dropped_but_children_survive(
        self,
    ) -> None:
        # Real CVC CCTP (Océania, Quaero, IFO_MAR) are full of narrative
        # sub-headings ("Généralités", "Principe", "Acoustique"...) that a
        # real economist never turns into a DPGF line — only their genuine
        # priceable children (if any) survive into the DPGF.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 12 CVC.docx"
            document = Document()
            document.add_heading("LOT 12 — CVC", level=1)
            document.add_heading("4.1 DESCRIPTION DES OUVRAGES DE CVC", level=2)
            document.add_heading("4.1.1 Généralités", level=3)
            document.add_paragraph("Le présent lot comprend la fourniture et la pose.")
            document.add_heading("4.1.2 Acoustique", level=3)
            document.add_paragraph("Les niveaux sonores respectent la réglementation.")
            document.add_heading("4.2 Production de chauffage", level=2)
            document.add_heading("4.2.1 Vase d'expansion", level=3)
            document.add_paragraph("Fourniture et pose, quantité 2.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_generic_titles")

        designations = {line["designation"] for line in lot["lines"]}
        self.assertNotIn("Généralités", designations)
        self.assertNotIn("Acoustique", designations)
        self.assertIn("Vase d'expansion", designations)

    def test_cvc_family_units_mined_from_real_dpgf(self) -> None:
        # Validated against 96 real CVC projects mined from
        # /Volumes/PARTAGE/ME/B_PROJETS (17 617 designation/unit pairs) — a
        # first pass based on only 3 CCTP/DPGF wrongly concluded "réseaux
        # aérauliques" → ml (real large-scale purity: 80 % Ens, not ml) and
        # found no real support for "passerelle"/"automate" (both dropped,
        # not replaced — see FAMILY_UNIT_OVERRIDES["cvc"] comment).
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 12 CVC.docx"
            document = Document()
            document.add_heading("LOT 12 — CVC", level=1)
            document.add_heading("4.1 DESCRIPTION DES OUVRAGES DE CVC", level=2)
            document.add_heading("4.1.1 Réseaux aérauliques", level=3)
            document.add_paragraph("Réseau de gaines, section rectangulaire et circulaire.")
            document.add_heading("4.1.2 Vase d'expansion", level=3)
            document.add_paragraph("Fourniture et pose, quantité 2.")
            document.add_heading("4.1.3 Caisson d'extraction", level=3)
            document.add_paragraph("Caisson avec équipements et accessoires.")
            document.add_heading("4.1.4 Boitier de répartition", level=3)
            document.add_paragraph("Boîtier maître, quantité 1.")
            document.add_heading("4.1.5 Soupape de sécurité", level=3)
            document.add_paragraph("Fourniture et pose, quantité 1.")
            document.add_heading("4.1.6 Désemboueur magnétique", level=3)
            document.add_paragraph("Complet, quantité 1.")
            document.add_heading("4.1.7 Régulation", level=3)
            document.add_paragraph("Régulation de l'installation.")
            document.add_heading("4.1.8 Collecteur départ retour", level=3)
            document.add_paragraph("Collecteur avec équipements.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_cvc_units")

        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Réseaux aérauliques"]["unit"], "Ens")
        self.assertEqual(by_designation["Vase d'expansion"]["unit"], "U")
        self.assertEqual(by_designation["Caisson d'extraction"]["unit"], "U")
        self.assertEqual(by_designation["Boitier de répartition"]["unit"], "U")
        self.assertEqual(by_designation["Soupape de sécurité"]["unit"], "U")
        self.assertEqual(by_designation["Désemboueur magnétique"]["unit"], "U")
        self.assertEqual(by_designation["Régulation"]["unit"], "Ens")
        self.assertEqual(by_designation["Collecteur départ retour"]["unit"], "Ens")

    def _lot_and_parse(
        self, directory: str, filename: str, lot_heading: str, items: list[tuple[str, str]]
    ) -> dict:
        path = Path(directory) / filename
        document = Document()
        document.add_heading(lot_heading, level=1)
        document.add_heading("1.1 DESCRIPTION DES OUVRAGES", level=2)
        for index, (title, paragraph) in enumerate(items, start=1):
            document.add_heading(f"1.1.{index} {title}", level=3)
            document.add_paragraph(paragraph)
        document.save(path)
        return parse_document(extract_document(path), f"src_{filename}")

    def test_str_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 88 projets Structure/Gros Œuvre depuis
        # /Volumes/PARTAGE/ME/B_PROJETS.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 03 GROS OEUVRE.docx",
                "LOT 03 — GROS ŒUVRE",
                [
                    ("Maçonnerie de remplissage", "Fourniture et pose, quantité 40."),
                    ("Panneau de signalisation de chantier", "Fourniture et pose, quantité 2."),
                    ("Branchements provisoires de chantier", "Ensemble, forfait."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Maçonnerie de remplissage"]["unit"], "m²")
        self.assertEqual(by_designation["Panneau de signalisation de chantier"]["unit"], "U")
        self.assertEqual(by_designation["Branchements provisoires de chantier"]["unit"], "Ens")

    def test_str_remblaiements_plural_now_covered(self) -> None:
        # "remblai" seul ne matchait pas ses formes dérivées ("remblaiements",
        # 97 % de pureté m³ sur 30 projets réels) — régression sur la règle
        # générique élargie, pas sur une FAMILY_UNIT_OVERRIDES.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 03 GROS OEUVRE.docx",
                "LOT 03 — GROS ŒUVRE",
                [("Remblaiements compactés", "Mise en œuvre, quantité 120.")],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Remblaiements compactés"]["unit"], "m³")

    def test_plomberie_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 60 projets Plomberie/Sanitaire, le plus riche des 12
        # oficios minés (120 candidats propres).
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 07 PLOMBERIE.docx",
                "LOT 07 — PLOMBERIE SANITAIRE",
                [
                    ("Dégorgement des canalisations", "Ensemble, forfait."),
                    ("Tampons de visite", "Fourniture et pose, quantité 3."),
                    ("Tube acier diam 20/27", "Section courante, quantité 15."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Dégorgement des canalisations"]["unit"], "Ens")
        self.assertEqual(by_designation["Tampons de visite"]["unit"], "Ens")
        self.assertEqual(by_designation["Tube acier diam 20/27"]["unit"], "ml")

    def test_serrurerie_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 67 projets Serrurerie/Métallerie.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 09 SERRURERIE.docx",
                "LOT 09 — SERRURERIE METALLERIE",
                [
                    ("Portillon métallique", "Fourniture et pose, quantité 1."),
                    ("Vantail de porte coulissante", "Fourniture et pose, quantité 2."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Portillon métallique"]["unit"], "U")
        self.assertEqual(by_designation["Vantail de porte coulissante"]["unit"], "U")

    def test_couverture_family_units_and_descente_family_conflict(self) -> None:
        # Minage réel : 65 projets Couverture/Étanchéité/Bardage. "descente"
        # penche vers U tous lots confondus, mais vers ml (80 %, 10 projets)
        # une fois qu'on sait que le lot est couverture/bardage (descentes
        # d'eaux pluviales facturées au mètre) — cas d'école pour
        # FAMILY_UNIT_OVERRIDES.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 04 COUVERTURE.docx",
                "LOT 04 — COUVERTURE ETANCHEITE BARDAGE",
                [
                    ("Couvertines aluminium", "Fourniture et pose, quantité 25."),
                    ("Descente eaux pluviales", "Fourniture et pose, quantité 12."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Couvertines aluminium"]["unit"], "ml")
        self.assertEqual(by_designation["Descente eaux pluviales"]["unit"], "ml")

    def test_menuiserie_exterieure_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 63 projets Menuiserie extérieure.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 10 MENUISERIE EXT.docx",
                "LOT 10 — MENUISERIE EXTERIEURE",
                [
                    ("Imposte fixe", "Fourniture et pose, quantité 4."),
                    ("Appuis de fenêtre aluminium", "Fourniture et pose, quantité 8."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Imposte fixe"]["unit"], "U")
        self.assertEqual(by_designation["Appuis de fenêtre aluminium"]["unit"], "ml")

    def test_menuiserie_interieure_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 39 projets Menuiserie intérieure.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 11 MENUISERIE INT.docx",
                "LOT 11 — MENUISERIE INTERIEURE",
                [
                    ("Signalétique intérieure", "Ensemble, forfait."),
                    ("Placard de rangement", "Fourniture et pose, quantité 2."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Signalétique intérieure"]["unit"], "Ens")
        self.assertEqual(by_designation["Placard de rangement"]["unit"], "U")

    def test_cloisons_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 58 projets Cloisons/Doublages/Plafonds.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 12 CLOISONS.docx",
                "LOT 12 — CLOISONS DOUBLAGES PLAFONDS",
                [
                    ("Parement hydrofuge", "Fourniture et pose, quantité 30."),
                    ("Ossature métallique", "Fourniture et pose, quantité 25."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Parement hydrofuge"]["unit"], "m²")
        self.assertEqual(by_designation["Ossature métallique"]["unit"], "m²")

    def test_peinture_family_units_and_beton_family_conflict(self) -> None:
        # Minage réel : 59 projets Peinture. "beton" penche vers m³ (un
        # volume coulé) tous lots confondus, mais vers m² (83 %, 12 projets)
        # en peinture — c'est le support à peindre, jamais le volume — même
        # cas d'école que "descente" en couverture ci-dessus.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 13 PEINTURE.docx",
                "LOT 13 — PEINTURE",
                [
                    ("Préparation des supports béton", "Ponçage et rebouchage, quantité 80."),
                    ("Lasure sur bois", "Deux couches, quantité 20."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Préparation des supports béton"]["unit"], "m²")
        self.assertEqual(by_designation["Lasure sur bois"]["unit"], "m²")

    def test_revetements_sols_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 60 projets Revêtements de sols.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 14 SOLS.docx",
                "LOT 14 — REVETEMENTS DE SOLS",
                [
                    ("Siphon de sol", "Fourniture et pose, quantité 3."),
                    ("Sous-couche acoustique", "Fourniture et pose, quantité 60."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Siphon de sol"]["unit"], "U")
        self.assertEqual(by_designation["Sous-couche acoustique"]["unit"], "m²")

    def test_espaces_verts_family_units_mined_from_real_dpgf(self) -> None:
        # Minage réel : 29 projets Espaces verts/Clôtures/Nettoyage.
        with tempfile.TemporaryDirectory() as directory:
            lot = self._lot_and_parse(
                directory,
                "CCTP LOT 15 ESPACES VERTS.docx",
                "LOT 15 — ESPACES VERTS",
                [
                    ("Arbres tiges", "Fourniture et pose, quantité 12."),
                    ("Paillage des massifs", "Mise en œuvre, quantité 50."),
                ],
            )
        by_designation = {line["designation"]: line for line in lot["lines"]}
        self.assertEqual(by_designation["Arbres tiges"]["unit"], "U")
        self.assertEqual(by_designation["Paillage des massifs"]["unit"], "m²")

    def test_lot_pattern_accepts_missing_separator_after_code(self) -> None:
        # Real IFO_MAR CCTP text: "E-206 - CCTP LOT 6 CVC - DESENFUMAGE" — no
        # dash/colon directly after the lot code, just a space before the
        # title ("CVC - DESENFUMAGE"); the dash there belongs to the title
        # itself, not to the code/title separator.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 06 CVC-DESENFUMAGE.docx"
            document = Document()
            document.add_paragraph("E-206 - CCTP LOT 6 CVC - DESENFUMAGE")
            document.add_heading("Description des ouvrages de chauffage", level=1)
            document.add_heading("2.1 Généralités", level=2)
            document.add_paragraph("Texte de généralités.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_lot_identity")

        self.assertEqual(lot["code"], "06")
        self.assertIn("CVC", lot["title"])
        self.assertIn("DESENFUMAGE", lot["title"])

    def test_decomposition_rule_adds_flagged_extra_lines(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 15 CVC.docx"
            document = Document()
            document.add_heading("LOT 15 — CVC", level=1)
            document.add_heading("15.1 DESCRIPTION DES OUVRAGES", level=2)
            document.add_heading("15.1.1 Tube cuivre", level=3)
            document.add_paragraph(
                "Fourniture et pose de tube cuivre pour l'ensemble des réseaux."
            )
            document.save(path)
            lot = parse_document(extract_document(path), "src_decomp")

        designations = {line["designation"]: line for line in lot["lines"]}
        self.assertIn("Tube cuivre diam. 12/14", designations)
        extra = designations["Tube cuivre diam. 12/14"]
        self.assertEqual(extra["origin"], "rule-derived")
        self.assertEqual(extra["review_status"], "to_review")
        self.assertIsNone(extra["quantity"])
        self.assertIn("implique", extra["review_reason"])
        # The CCTP-derived line itself is untouched.
        self.assertEqual(designations["Tube cuivre"]["origin"], "deterministic-v2")

    def test_second_wave_decomposition_rules_fire(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 16 PLOMBERIE.docx"
            document = Document()
            document.add_heading("LOT 16 — PLOMBERIE", level=1)
            document.add_heading("Description des ouvrages", level=1)
            document.add_heading("Stérilisation du réseau", level=2)
            document.add_paragraph("Stérilisation complète du réseau d'eau potable.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_second_wave")

        designations = {line["designation"] for line in lot["lines"]}
        self.assertIn("Repérage étiquetage", designations)
        self.assertIn("Analyse d'eau", designations)

    def _build_control_trigger_docx(self, path: Path, lot_heading: str) -> None:
        document = Document()
        document.add_heading(lot_heading, level=1)
        document.add_heading("3.4 OUVRAGES D'ASSAINISSEMENT", level=2)
        document.add_heading("3.4.1 Ouvrages d'assainissement", level=3)
        document.add_paragraph("Ensemble des canalisations et regards, au forfait.")
        document.add_heading("3.4.2 Nettoyage de fin de chantier", level=3)
        document.add_paragraph("Ensemble, forfait.")
        document.add_heading("3.5 REFECTION DE VOIRIE", level=2)
        document.add_heading("3.5.1 Réglage de fond de forme", level=3)
        document.add_paragraph("Surface en m².")
        document.save(path)

    def test_vrd_defers_rule_derived_control_lines_to_end_of_section(self) -> None:
        # Real DPGF group "Contrôle qualité" lines at the end of their own
        # section (e.g. right before "Sous-total 3.4"), not wherever the
        # triggering CCTP text happened to sit — validated against a real
        # Bonduelle LOT 01 VRD CCTP/DPGF pair.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 01 VRD.docx"
            self._build_control_trigger_docx(path, "LOT 01 — VRD")
            lot = parse_document(extract_document(path), "src_vrd_control")

        designations = [line["designation"] for line in lot["lines"]]
        control_index = designations.index(
            "Contrôle qualité des ouvrages d'assainissement (étanchéité RV)"
        )
        cleanup_index = designations.index("Nettoyage de fin de chantier")
        next_section_index = designations.index("REFECTION DE VOIRIE")
        self.assertGreater(control_index, cleanup_index)
        self.assertLess(control_index, next_section_index)

    def test_non_vrd_lot_keeps_control_lines_in_original_position(self) -> None:
        # The end-of-section grouping is a VRD DPGF convention, not a
        # universal rule — other trades keep the historical in-place
        # insertion until a similar convention is confirmed for them too.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 03 GROS OEUVRE.docx"
            self._build_control_trigger_docx(path, "LOT 03 — GROS ŒUVRE")
            lot = parse_document(extract_document(path), "src_gros_oeuvre_control")

        designations = [line["designation"] for line in lot["lines"]]
        control_index = designations.index(
            "Contrôle qualité des ouvrages d'assainissement (étanchéité RV)"
        )
        cleanup_index = designations.index("Nettoyage de fin de chantier")
        self.assertLess(control_index, cleanup_index)

    def test_decomposition_rule_skips_variant_already_written_in_cctp(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 15 CVC.docx"
            document = Document()
            document.add_heading("LOT 15 — CVC", level=1)
            document.add_heading("15.1 DESCRIPTION DES OUVRAGES", level=2)
            document.add_heading("15.1.1 Tube cuivre", level=3)
            document.add_paragraph("Fourniture et pose de tube cuivre.")
            document.add_heading("15.1.2 Tube cuivre diam. 12/14", level=3)
            document.add_paragraph("Section courante en diamètre 12/14, quantité 30 ml.")
            document.save(path)
            lot = parse_document(extract_document(path), "src_decomp_dup")

        matches = [
            line
            for line in lot["lines"]
            if line["designation"] == "Tube cuivre diam. 12/14"
        ]
        # Written explicitly in the CCTP: the rule must not duplicate it.
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0]["origin"], "deterministic-v2")

    def test_lot_family_breaks_ties_when_codes_collide(self) -> None:
        self.assertEqual(classify_lot_family("LOT 03 — ELECTRICITE CFO/CFA")[0], "electricite")
        self.assertEqual(classify_lot_family("LOT 03 — GROS OEUVRE")[0], "fondations_gros_oeuvre")
        electricite = {"code": "03", "title": "LOT 03 — ELECTRICITE CFO/CFA"}
        gros_oeuvre = {"code": "03", "title": "LOT 03 — GROS OEUVRE"}
        # Same code from two independently numbered CCTP: gros œuvre still
        # comes before the technical lots in the merged workbook.
        ordered = sorted([electricite, gros_oeuvre], key=lot_sort_key)
        self.assertEqual([lot["title"] for lot in ordered], [gros_oeuvre["title"], electricite["title"]])
        # A project with its own coherent, distinct numbering is untouched.
        self.assertEqual(
            sorted(
                [
                    {"code": "12", "title": "LOT 12 — ELECTRICITE"},
                    {"code": "03", "title": "LOT 03 — GROS OEUVRE"},
                ],
                key=lot_sort_key,
            )[0]["code"],
            "03",
        )


class ExcelTests(unittest.TestCase):
    def test_export_matches_model_and_uses_correct_totals(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "CCTP LOT 03.docx"
            make_docx(source)
            lot = parse_document(extract_document(source), "src_export")
            manual = {
                "id": "manual_1",
                "kind": "item",
                "level": 3,
                "code": "3.2.2",
                "designation": "Objet ajouté manuellement",
                "description": "",
                "unit": "U",
                "quantity": 2,
                "unit_price": None,
                "included": True,
                "confidence": 1,
                "review_status": "validated",
                "review_reason": "",
                "source_id": "src_export",
                "source_page": None,
                "source_excerpt": "Ajout manuel",
                "origin": "manual",
            }
            lot["lines"].append(manual)
            output = Path(directory) / "DPGF.xlsx"
            export_analysis(sample_analysis(lot), output)
            with ZipFile(output) as archive:
                logo_members = [
                    name
                    for name in archive.namelist()
                    if name.startswith("xl/media/")
                ]
                self.assertEqual(len(logo_members), 1)
                self.assertTrue(logo_members[0].endswith(".png"))
                self.assertTrue(
                    archive.read(logo_members[0]).startswith(
                        b"\x89PNG\r\n\x1a\n"
                    )
                )
            workbook = load_workbook(output, data_only=False)
            sheet = workbook.active

        self.assertIn("LOT 03", sheet["B8"].value)
        self.assertEqual(len(sheet._images), 1)
        self.assertEqual(sheet.row_dimensions[2].height, 51)
        values = {
            str(sheet.cell(row, 3).value): row
            for row in range(1, sheet.max_row + 1)
            if sheet.cell(row, 3).value
        }
        manual_row = values["Objet ajouté manuellement"]
        self.assertEqual(sheet.cell(manual_row, 7).value, f'=IF(OR(E{manual_row}="",F{manual_row}=""),"",E{manual_row}*F{manual_row})')
        self.assertFalse(sheet.cell(manual_row, 6).protection.locked)
        self.assertFalse(sheet.cell(manual_row, 5).protection.locked)
        self.assertIsNotNone(sheet.cell(manual_row, 4).comment)
        vat_row = next(row for label, row in values.items() if label.startswith("TVA à"))
        total_row = values["TOTAL € T.T.C."]
        self.assertIn(f"G{vat_row - 1}*20", sheet.cell(vat_row, 7).value)
        self.assertIn(f"G{vat_row - 1}", sheet.cell(total_row, 7).value)
        self.assertFalse(sheet.protection.sheet)

    def test_export_only_subtotals_x_x_and_styles_deeper_objects(self) -> None:
        def line(
            code: str,
            designation: str,
            kind: str,
            level: int,
            *,
            included: bool = True,
        ) -> dict:
            return {
                "id": f"line_{code}",
                "kind": kind,
                "level": level,
                "code": code,
                "designation": designation,
                "description": "",
                "unit": None if kind == "section" else "U",
                "quantity": None,
                "unit_price": None,
                "included": included,
                "confidence": 1,
                "review_status": "validated",
                "review_reason": "",
                "source_id": "src_tree",
                "source_page": 1,
                "source_excerpt": designation,
                "origin": "deterministic-v2",
                "unit_source": None if kind == "section" else "rule",
            }

        lot = {
            "id": "lot_tree",
            "code": "05",
            "title": "Menuiseries extérieures",
            "source_id": "src_tree",
            "warnings": [],
            "lines": [
                line("4", "DESCRIPTION DES OUVRAGES", "section", 1),
                line("4.5", "MUR RIDEAU ALUMINIUM", "section", 2),
                line("4.5.1", "Dépose mur rideau", "item", 3),
                line("4.5.2", "Remplacement vitrages", "item", 3),
                line("4.6", "Porte vitrée automatique", "item", 2),
                line("4.7", "ENSEMBLE TECHNIQUE", "section", 2),
                line("4.7.1", "SOUS-ENSEMBLE", "section", 3),
                line("4.7.1.1", "Composant", "item", 4),
                line("4.7.2", "Variante", "item", 3, included=False),
            ],
        }
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "tree.xlsx"
            export_analysis(sample_analysis(lot), output)
            sheet = load_workbook(output, data_only=False).active

        rows_by_code = {
            str(sheet.cell(row, 2).value): row
            for row in range(12, sheet.max_row + 1)
            if sheet.cell(row, 2).value
        }
        rows_by_label = {
            str(sheet.cell(row, 3).value): row
            for row in range(12, sheet.max_row + 1)
            if sheet.cell(row, 3).value
        }
        self.assertLess(rows_by_label["Sous-total 4.5"], rows_by_code["4.6"])
        self.assertEqual(sheet.cell(rows_by_code["4.6"], 4).value, "U")
        self.assertTrue(sheet.cell(rows_by_code["4.6"], 2).font.bold)
        self.assertTrue(sheet.cell(rows_by_code["4.6"], 3).font.bold)
        self.assertLess(rows_by_code["4.6"], rows_by_label["Sous-total 4.6"])
        self.assertLess(rows_by_label["Sous-total 4.6"], rows_by_code["4.7"])
        self.assertEqual(
            sheet.cell(rows_by_label["Sous-total 4.6"], 7).value,
            f"=G{rows_by_code['4.6']}",
        )
        self.assertEqual(
            sheet.cell(rows_by_code["4.6"], 2).style_id,
            sheet.cell(rows_by_code["4.5"], 2).style_id,
        )
        self.assertFalse(sheet.cell(rows_by_code["4.5.1"], 3).font.bold)
        self.assertTrue(sheet.cell(rows_by_code["4.5.1"], 3).font.italic)
        self.assertEqual(sheet.cell(rows_by_code["4.5.1"], 3).font.sz, 8)
        self.assertEqual(sheet.cell(rows_by_code["4.5"], 3).font.sz, 9)
        self.assertEqual(sheet.cell(rows_by_code["4"], 3).font.sz, 9)
        self.assertNotIn("Sous-total 4.7.1", rows_by_label)
        self.assertTrue(sheet.cell(rows_by_code["4.7.1"], 3).font.italic)
        self.assertFalse(sheet.cell(rows_by_code["4.7.1"], 3).font.bold)
        self.assertFalse(sheet.cell(rows_by_code["4.7.2"], 2).font.bold)
        self.assertTrue(sheet.cell(rows_by_code["4.7.2"], 3).font.italic)
        self.assertLess(rows_by_code["4.7.2"], rows_by_label["Sous-total 4.7"])
        self.assertNotIn(
            f"G{rows_by_code['4.7.2']}",
            str(sheet.cell(rows_by_label["Sous-total 4.7"], 7).value),
        )
        total_row = rows_by_label["TOTAL € H.T. hors prorata"]
        self.assertNotIn(
            f"G{rows_by_label['Sous-total 4.5']}",
            str(sheet.cell(total_row, 7).value),
        )
        self.assertEqual(sheet.sheet_view.topLeftCell, "A1")
        self.assertFalse(any(sheet.row_dimensions[row].hidden for row in range(1, 9)))

    def test_export_creates_one_workbook_per_lot(self) -> None:
        def lot(code: str, title: str) -> dict:
            return {
                "id": f"lot_{code}",
                "code": code,
                "title": title,
                "source_id": f"src_{code}",
                "warnings": [],
                "lines": [
                    {
                        "id": f"line_{code}",
                        "kind": "item",
                        "level": 2,
                        "code": f"{int(code)}.1",
                        "designation": f"Objet du lot {code}",
                        "description": "",
                        "unit": "U",
                        "quantity": None,
                        "unit_price": None,
                        "included": True,
                        "confidence": 1,
                        "review_status": "validated",
                        "review_reason": "",
                        "source_id": f"src_{code}",
                        "source_page": 1,
                        "source_excerpt": title,
                        "origin": "deterministic-v2",
                        "unit_source": "rule",
                    }
                ],
            }

        lots = [lot("03", "Gros œuvre"), lot("04", "Menuiseries")]
        analysis = {
            **sample_analysis(lots[0]),
            "lots": lots,
            "stats": recompute_stats(lots),
        }
        with tempfile.TemporaryDirectory() as directory:
            outputs = export_analysis_files(analysis, Path(directory))
            labels = [
                load_workbook(output, data_only=False).active["B8"].value
                for output in outputs
            ]

        self.assertEqual(len(outputs), 2)
        self.assertEqual(len({output.name for output in outputs}), 2)
        self.assertIn("LOT 03", labels[0])
        self.assertIn("LOT 04", labels[1])


class ApiFlowTests(unittest.TestCase):
    def test_complete_history_edit_export_flow(self) -> None:
        client = TestClient(app)
        buffer = BytesIO()
        document = Document()
        document.add_heading("LOT 05 — MENUISERIES", level=1)
        document.add_heading("5.1 PORTES", level=2)
        document.add_heading("5.1.1 Bloc-porte intérieur", level=3)
        document.add_paragraph("Fourniture et pose à l'unité.")
        document.save(buffer)
        second_buffer = BytesIO()
        second_document = Document()
        second_document.add_heading("LOT 06 — PEINTURE", level=1)
        second_document.add_heading("6.1 PEINTURES INTÉRIEURES", level=2)
        second_document.add_heading("6.1.1 Peinture murale", level=3)
        second_document.add_paragraph("Préparation et peinture, unité de règlement : m².")
        second_document.save(second_buffer)
        response = client.post(
            "/api/v1/analyses",
            data={
                "project_name": "Projet test DPGF",
                "project_reference": "TEST-001",
                "phase": "DCE",
            },
            files=[
                (
                    "files",
                    (
                        "CCTP LOT 05.docx",
                        buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "CCTP LOT 06.docx",
                        second_buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(response.status_code, 202, response.text)
        analysis_id = response.json()["id"]
        try:
            detail = client.get(f"/api/v1/analyses/{analysis_id}")
            self.assertEqual(detail.status_code, 200, detail.text)
            payload = detail.json()
            self.assertIn(payload["status"], {"ready", "needs_review"})
            lot = payload["lots"][0]
            lot["lines"].append(
                {
                    "id": "manual_api",
                    "kind": "item",
                    "level": 3,
                    "code": "5.1.2",
                    "designation": "Objet manuel API",
                    "unit": "U",
                    "quantity": 4,
                    "included": True,
                    "review_status": "validated",
                    "origin": "manual",
                }
            )
            saved = client.put(
                f"/api/v1/analyses/{analysis_id}",
                json={"project": payload["project"], "lots": payload["lots"]},
            )
            self.assertEqual(saved.status_code, 200, saved.text)
            self.assertTrue(
                any(
                    line["designation"] == "Objet manuel API"
                    for line in saved.json()["lots"][0]["lines"]
                )
            )
            reprocessed = client.post(
                f"/api/v1/analyses/{analysis_id}/reprocess"
            )
            self.assertEqual(reprocessed.status_code, 202, reprocessed.text)
            refreshed = client.get(f"/api/v1/analyses/{analysis_id}").json()
            self.assertTrue(
                any(
                    line["designation"] == "Objet manuel API"
                    for line in refreshed["lots"][0]["lines"]
                )
            )
            revisions = list(
                (store.analysis_directory(analysis_id) / "revisions").glob("*.json")
            )
            self.assertEqual(len(revisions), 1)
            exported = client.post(f"/api/v1/analyses/{analysis_id}/export")
            self.assertEqual(exported.status_code, 200, exported.text)
            download = client.get(exported.json()["download_url"])
            self.assertEqual(download.status_code, 200)
            self.assertEqual(download.headers["content-type"], "application/zip")
            with ZipFile(BytesIO(download.content)) as archive:
                workbook_names = [
                    name for name in archive.namelist() if name.endswith(".xlsx")
                ]
                workbooks = [
                    load_workbook(
                        BytesIO(archive.read(name)),
                        data_only=False,
                    )
                    for name in workbook_names
                ]
            self.assertEqual(len(workbook_names), 2)
            self.assertTrue(
                any(
                    cell.value == "Objet manuel API"
                    for workbook in workbooks
                    for row in workbook.active.iter_rows()
                    for cell in row
                )
            )
            history = client.get("/api/v1/analyses")
            self.assertTrue(any(item["id"] == analysis_id for item in history.json()["analyses"]))
            tco = client.get(f"/api/v1/analyses/{analysis_id}/tco")
            self.assertEqual(tco.json()["schema_version"], "1.0")
        finally:
            client.delete(f"/api/v1/analyses/{analysis_id}")

    def test_pm_unit_is_accepted_and_preserved(self) -> None:
        # "PM" (Pour Mémoire) is a real unit in real DPGF (Quaero, IFO_MAR)
        # for provisional/informational lines — save_analysis must not
        # silently coerce it back to "Ens" like it does for unknown units.
        client = TestClient(app)
        buffer = BytesIO()
        document = Document()
        document.add_heading("LOT 12 — CVC", level=1)
        document.add_heading("4.1 DESCRIPTION DES OUVRAGES DE CVC", level=2)
        document.add_heading("4.1.1 Base de calcul", level=3)
        document.add_paragraph("Bilan calorifique et frigorifique.")
        document.save(buffer)
        response = client.post(
            "/api/v1/analyses",
            data={"project_name": "Projet PM", "project_reference": "PM-001", "phase": "DCE"},
            files=[
                (
                    "files",
                    (
                        "CCTP LOT 12.docx",
                        buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(response.status_code, 202, response.text)
        analysis_id = response.json()["id"]
        try:
            payload = client.get(f"/api/v1/analyses/{analysis_id}").json()
            lot = payload["lots"][0]
            lot["lines"].append(
                {
                    "id": "manual_pm",
                    "kind": "item",
                    "level": 3,
                    "code": "4.1.2",
                    "designation": "Bilan calorifique et frigorifique",
                    "unit": "PM",
                    "quantity": None,
                    "included": True,
                    "review_status": "validated",
                    "origin": "manual",
                }
            )
            saved = client.put(
                f"/api/v1/analyses/{analysis_id}",
                json={"project": payload["project"], "lots": payload["lots"]},
            )
            self.assertEqual(saved.status_code, 200, saved.text)
            saved_line = next(
                line
                for line in saved.json()["lots"][0]["lines"]
                if line["id"] == "manual_pm"
            )
            self.assertEqual(saved_line["unit"], "PM")
        finally:
            client.delete(f"/api/v1/analyses/{analysis_id}")


class MultiUserPermissionTests(unittest.TestCase):
    """DPGF Résumé CCTP is opt-in shared: an analysis is only visible to its
    owner or to someone it was explicitly shared with (view or edit), plus
    whoever already holds edit rights by role (Admin/Copil/superuser) can
    still open it directly even without a share row."""

    def _identity(
        self, sub: str, role: str, *, is_superuser: bool = False
    ) -> auth.UserIdentity:
        return auth.UserIdentity(
            sub=sub,
            name=f"Utilisateur {sub}",
            email=f"{sub}@example.com",
            username=sub,
            role=role,
            is_superuser=is_superuser,
        )

    def _create_as(self, client: TestClient, identity: auth.UserIdentity) -> str:
        buffer = BytesIO()
        document = Document()
        document.add_heading("LOT 05 — MENUISERIES", level=1)
        document.add_heading("5.1 PORTES", level=2)
        document.add_heading("5.1.1 Bloc-porte intérieur", level=3)
        document.add_paragraph("Fourniture et pose à l'unité.")
        document.save(buffer)
        with patch.object(auth.service, "current_user", return_value=identity):
            response = client.post(
                "/api/v1/analyses",
                data={
                    "project_name": "Projet partagé",
                    "project_reference": "SHARE-001",
                    "phase": "DCE",
                },
                files=[
                    (
                        "files",
                        (
                            "CCTP LOT 05.docx",
                            buffer.getvalue(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                ],
            )
        self.assertEqual(response.status_code, 202, response.text)
        return response.json()["id"]

    def test_non_owner_without_share_is_forbidden_and_hidden(self) -> None:
        client = TestClient(app)
        owner = self._identity("user-a", "Collaborateur")
        other = self._identity("user-b", "Collaborateur")
        analysis_id = self._create_as(client, owner)
        try:
            with patch.object(auth.service, "current_user", return_value=other):
                detail = client.get(f"/api/v1/analyses/{analysis_id}")
                self.assertEqual(detail.status_code, 403)

                history = client.get("/api/v1/analyses")
                self.assertEqual(history.status_code, 200)
                ids = [item["id"] for item in history.json()["analyses"]]
                self.assertNotIn(analysis_id, ids)
        finally:
            with patch.object(auth.service, "current_user", return_value=owner):
                client.delete(f"/api/v1/analyses/{analysis_id}")

    def test_view_share_grants_read_only_access(self) -> None:
        client = TestClient(app)
        owner = self._identity("user-a", "Collaborateur")
        viewer = self._identity("user-b", "Collaborateur")
        analysis_id = self._create_as(client, owner)
        try:
            with patch.object(auth.service, "current_user", return_value=owner):
                share_response = client.put(
                    f"/api/v1/analyses/{analysis_id}/shares",
                    json={"shares": [{"email": "user-b@example.com", "permission": "view"}]},
                )
                self.assertEqual(share_response.status_code, 200, share_response.text)

            with patch.object(auth.service, "current_user", return_value=viewer):
                detail = client.get(f"/api/v1/analyses/{analysis_id}")
                self.assertEqual(detail.status_code, 200, detail.text)
                payload = detail.json()
                self.assertFalse(payload["can_edit"])
                self.assertFalse(payload["can_manage_sharing"])
                self.assertEqual(payload["shares"], [])  # hidden from non-managers

                history = client.get("/api/v1/analyses")
                ids = [item["id"] for item in history.json()["analyses"]]
                self.assertIn(analysis_id, ids)

                for response in (
                    client.put(
                        f"/api/v1/analyses/{analysis_id}",
                        json={"project": payload["project"], "lots": payload["lots"]},
                    ),
                    client.post(f"/api/v1/analyses/{analysis_id}/reprocess"),
                    client.post(f"/api/v1/analyses/{analysis_id}/export"),
                    client.delete(f"/api/v1/analyses/{analysis_id}"),
                ):
                    self.assertEqual(response.status_code, 403)
        finally:
            with patch.object(auth.service, "current_user", return_value=owner):
                client.delete(f"/api/v1/analyses/{analysis_id}")

    def test_edit_share_can_edit_but_not_manage_sharing(self) -> None:
        client = TestClient(app)
        owner = self._identity("user-a", "Collaborateur")
        editor = self._identity("user-c", "Collaborateur")
        analysis_id = self._create_as(client, owner)
        try:
            with patch.object(auth.service, "current_user", return_value=owner):
                client.put(
                    f"/api/v1/analyses/{analysis_id}/shares",
                    json={"shares": [{"email": "user-c@example.com", "permission": "edit"}]},
                )

            with patch.object(auth.service, "current_user", return_value=editor):
                detail = client.get(f"/api/v1/analyses/{analysis_id}")
                payload = detail.json()
                self.assertTrue(payload["can_edit"])
                self.assertFalse(payload["can_manage_sharing"])

                payload["lots"][0]["lines"].append(
                    {
                        "id": "manual_editor",
                        "kind": "item",
                        "level": 3,
                        "code": "5.1.2",
                        "designation": "Objet ajouté par un éditeur partagé",
                        "unit": "U",
                        "quantity": 2,
                        "included": True,
                        "review_status": "validated",
                        "origin": "manual",
                    }
                )
                saved = client.put(
                    f"/api/v1/analyses/{analysis_id}",
                    json={"project": payload["project"], "lots": payload["lots"]},
                )
                self.assertEqual(saved.status_code, 200, saved.text)
                # Regression: an editor who only has a share must not become
                # the owner (app/store.py update_analysis always preserves
                # the original owner).
                self.assertEqual(saved.json()["owner"]["sub"], "user-a")

                # An "edit" share does not let you re-share with others —
                # sharing stays reserved for the owner/Admin/Copil.
                escalation = client.put(
                    f"/api/v1/analyses/{analysis_id}/shares",
                    json={"shares": [{"email": "user-d@example.com", "permission": "view"}]},
                )
                self.assertEqual(escalation.status_code, 403)
        finally:
            with patch.object(auth.service, "current_user", return_value=owner):
                client.delete(f"/api/v1/analyses/{analysis_id}")

    def test_admin_role_can_open_by_id_but_not_listed_without_share(self) -> None:
        client = TestClient(app)
        owner = self._identity("user-a", "Collaborateur")
        admin = self._identity("user-e", "Admin")
        analysis_id = self._create_as(client, owner)
        try:
            with patch.object(auth.service, "current_user", return_value=admin):
                detail = client.get(f"/api/v1/analyses/{analysis_id}")
                self.assertEqual(detail.status_code, 200, detail.text)
                self.assertTrue(detail.json()["can_edit"])

                history = client.get("/api/v1/analyses")
                ids = [item["id"] for item in history.json()["analyses"]]
                self.assertNotIn(analysis_id, ids)

                delete_response = client.delete(f"/api/v1/analyses/{analysis_id}")
                self.assertEqual(delete_response.status_code, 204, delete_response.text)
        except Exception:
            with patch.object(auth.service, "current_user", return_value=owner):
                client.delete(f"/api/v1/analyses/{analysis_id}")
            raise

    def test_replace_shares_rejects_invalid_permission(self) -> None:
        client = TestClient(app)
        owner = self._identity("user-a", "Collaborateur")
        analysis_id = self._create_as(client, owner)
        try:
            with patch.object(auth.service, "current_user", return_value=owner):
                response = client.put(
                    f"/api/v1/analyses/{analysis_id}/shares",
                    json={"shares": [{"email": "user-b@example.com", "permission": "admin"}]},
                )
                self.assertEqual(response.status_code, 422)
        finally:
            with patch.object(auth.service, "current_user", return_value=owner):
                client.delete(f"/api/v1/analyses/{analysis_id}")

    def test_directory_users_lists_mocked_group_members(self) -> None:
        client = TestClient(app)
        user = self._identity("user-a", "Collaborateur")
        members = [
            {"email": "user-b@example.com", "name": "Bea B", "username": "user-b"},
            {"email": "user-c@example.com", "name": "Cy C", "username": "user-c"},
        ]
        with patch.object(auth.service, "current_user", return_value=user), patch.object(
            directory_module, "list_group_members", return_value=members
        ):
            response = client.get("/api/v1/directory/users")
            self.assertEqual(response.status_code, 200, response.text)
            self.assertEqual(response.json()["users"], members)

            filtered = client.get("/api/v1/directory/users?q=bea")
            self.assertEqual(len(filtered.json()["users"]), 1)
            self.assertEqual(filtered.json()["users"][0]["email"], "user-b@example.com")


class LlmAssistTests(unittest.TestCase):
    """LIHA is opt-in (config.USE_LLM=False for the whole suite, see top of
    this file) and none of these tests hit the network: app.llm functions are
    mocked directly."""

    def test_perimeter_assist_recovers_from_wrong_dominant_chapter(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 08 COUVERTURE.docx"
            document = Document()
            document.add_heading("LOT 08 — COUVERTURE", level=1)
            # Chapter 2 (administrative) has 3 leaves, chapter 3 (the real
            # priceable works) has only 2 — the dominant-chapter heuristic
            # picks chapter 2 by sheer count, which is wrong.
            document.add_heading("2. GENERALITES", level=1)
            document.add_heading("2.1 Documents de référence", level=2)
            document.add_paragraph("Liste des documents applicables.")
            document.add_heading("2.2 Normes", level=2)
            document.add_paragraph("Liste des normes applicables.")
            document.add_heading("2.3 Prescriptions communes", level=2)
            document.add_paragraph("Prescriptions générales du marché.")
            document.add_heading("3. TRAVAUX DE COUVERTURE", level=1)
            document.add_heading("3.1 Ragréage", level=2)
            document.add_paragraph("Fourniture et mise en oeuvre d'un ragréage.")
            document.add_heading("3.2 Carrelage", level=2)
            document.add_paragraph("Fourniture et pose d'un carrelage.")
            document.save(path)
            extracted = extract_document(path)
            lot = parse_document(extracted, "src_perimeter")

        self.assertEqual(lot["perimeter"]["method"], "dominant_numbered_chapter")
        self.assertEqual(lot["perimeter"]["anchor_code"], "2")

        with patch.object(
            llm_module, "suggest_perimeter_anchor", return_value="3"
        ) as mock_suggest:
            new_lot, used = _apply_perimeter_assist(extracted, "src_perimeter", lot)

        mock_suggest.assert_called_once()
        self.assertTrue(used)
        self.assertEqual(new_lot["perimeter"]["method"], "llm_confirmed_anchor")
        designations = {line["designation"] for line in new_lot["lines"]}
        self.assertIn("Ragréage", designations)
        self.assertNotIn("Documents de référence", designations)

    def test_perimeter_assist_skips_already_strong_perimeter(self) -> None:
        lot = {"perimeter": {"method": "explicit_anchor"}, "lines": [], "title": "x"}
        with patch.object(llm_module, "suggest_perimeter_anchor") as mock_suggest:
            new_lot, used = _apply_perimeter_assist(object(), "src", lot)
        mock_suggest.assert_not_called()
        self.assertFalse(used)
        self.assertIs(new_lot, lot)

    def test_perimeter_assist_rejects_anchor_that_yields_no_items(self) -> None:
        # The confirmed anchor is a real candidate but happens to have no
        # further numbered breakdown in this document — the forced re-parse
        # comes back with zero priceable items. Keep the original result
        # instead of silently replacing it with an empty one.
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "CCTP LOT 09 MENUISERIES.docx"
            document = Document()
            document.add_heading("LOT 09 — MENUISERIES", level=1)
            document.add_heading("2. GENERALITES", level=1)
            document.add_heading("2.1 Documents de référence", level=2)
            document.add_paragraph("Liste des documents.")
            document.add_heading("2.2 Normes", level=2)
            document.add_paragraph("Liste des normes.")
            document.add_heading("3. TRAVAUX", level=1)
            document.add_paragraph("Paragraphe sans sous-titre numéroté.")
            document.save(path)
            extracted = extract_document(path)
            lot = parse_document(extracted, "src_empty_anchor")

        self.assertEqual(lot["perimeter"]["method"], "dominant_numbered_chapter")
        original_items = sum(1 for line in lot["lines"] if line["kind"] == "item")

        with patch.object(llm_module, "suggest_perimeter_anchor", return_value="3"):
            new_lot, used = _apply_perimeter_assist(extracted, "src_empty_anchor", lot)

        self.assertFalse(used)
        self.assertIs(new_lot, lot)
        self.assertEqual(
            sum(1 for line in new_lot["lines"] if line["kind"] == "item"), original_items
        )

    def test_unit_assist_fills_default_without_touching_rule_based_units(self) -> None:
        lot = {
            "lines": [
                {
                    "id": "a",
                    "kind": "item",
                    "designation": "Robinet de puisage",
                    "unit": "U",
                    "unit_source": "rule",
                },
                {
                    "id": "b",
                    "kind": "item",
                    "designation": "Traitement de surface spécifique",
                    "unit": "Ens",
                    "unit_source": "default",
                },
            ]
        }
        with patch.object(
            llm_module, "suggest_units", return_value={"b": "m²"}
        ) as mock_suggest:
            used = _apply_unit_assist(lot)

        mock_suggest.assert_called_once_with(
            [{"code": "b", "designation": "Traitement de surface spécifique"}]
        )
        self.assertTrue(used)
        by_id = {line["id"]: line for line in lot["lines"]}
        self.assertEqual(by_id["a"]["unit"], "U")
        self.assertEqual(by_id["a"]["unit_source"], "rule")
        self.assertEqual(by_id["b"]["unit"], "m²")
        self.assertEqual(by_id["b"]["unit_source"], "llm")
        self.assertEqual(by_id["b"]["unit_confidence"], 0.7)

    def test_suggest_perimeter_anchor_ignores_code_outside_candidates(self) -> None:
        with patch.object(llm_module, "available", return_value=True), patch.object(
            llm_module, "_chat_completion", return_value='{"anchor_code": "99"}'
        ):
            result = llm_module.suggest_perimeter_anchor(
                [{"code": "3", "title": "Travaux", "level": 1}], "LOT 08"
            )
        self.assertIsNone(result)

    def test_suggest_units_ignores_invalid_unit_and_unknown_code(self) -> None:
        with patch.object(llm_module, "available", return_value=True), patch.object(
            llm_module,
            "_chat_completion",
            return_value='{"units": {"a": "litres", "zzz": "U"}}',
        ):
            result = llm_module.suggest_units([{"code": "a", "designation": "Foo"}])
        self.assertEqual(result, {})

    def test_suggest_units_accepts_pm(self) -> None:
        # "PM" (Pour Mémoire) is a real DPGF unit for provisional/
        # informational lines (Quaero, IFO_MAR) — must not be rejected like
        # a genuinely invalid unit ("litres" above).
        with patch.object(llm_module, "available", return_value=True), patch.object(
            llm_module,
            "_chat_completion",
            return_value='{"units": {"a": "PM"}}',
        ):
            result = llm_module.suggest_units([{"code": "a", "designation": "Base de calcul"}])
        self.assertEqual(result, {"a": "PM"})


if __name__ == "__main__":
    unittest.main()
