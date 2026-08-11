from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from statistics import median
from typing import Iterable

import fitz
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class TextBlock:
    text: str
    page: int | None
    order: int
    style: str = ""
    font_size: float = 0.0
    bold: bool = False
    source_kind: str = "paragraph"
    x0: float | None = None
    y0: float | None = None
    x1: float | None = None
    y1: float | None = None
    page_width: float | None = None
    page_height: float | None = None

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class ExtractedDocument:
    path: Path
    blocks: list[TextBlock]
    page_count: int
    character_count: int
    warnings: list[str]

    @property
    def text(self) -> str:
        return "\n".join(block.text for block in self.blocks)


_CODE_ONLY = re.compile(
    r"^\s*(?:ARTICLE\s+)?[A-Z]?\d{1,3}(?:[.\-]\d{1,3}){0,6}[.)]?\s*$",
    re.IGNORECASE,
)


def _clean_text(value: str) -> str:
    value = (
        value.replace("\u00a0", " ")
        .replace("\u200b", "")
        .replace("\uf0a7", "•")
    )
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def _repeat_key(value: str) -> str:
    value = value.casefold()
    value = re.sub(r"\b\d+\s*/\s*\d+\b", "#/#", value)
    value = re.sub(r"\bpage\s+\d+\b", "page #", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip(" -–—")


def _merge_split_headings(lines: list[dict]) -> list[dict]:
    """Recompose headings whose number and title are separate PDF text boxes."""
    merged: list[dict] = []
    index = 0
    while index < len(lines):
        current = lines[index]
        if _CODE_ONLY.match(current["text"]) and index + 1 < len(lines):
            following = lines[index + 1]
            same_baseline = abs(float(current["y0"]) - float(following["y0"])) <= max(
                2.6, float(current["size"]) * 0.22
            )
            to_the_right = float(following["x0"]) >= float(current["x0"]) + 8
            plausible_title = (
                2 <= len(following["text"]) <= 220
                and not _CODE_ONLY.match(following["text"])
                and not following["text"].startswith(("•", "-", "–"))
            )
            similar_style = (
                bool(current["bold"]) == bool(following["bold"])
                or abs(float(current["size"]) - float(following["size"])) <= 1.2
            )
            if same_baseline and to_the_right and plausible_title and similar_style:
                code = current["text"].strip()
                if code.endswith(".") and "." in code[:-1]:
                    code = code[:-1]
                current = {
                    **current,
                    "text": f"{code} {following['text']}",
                    "x1": max(float(current["x1"]), float(following["x1"])),
                    "y1": max(float(current["y1"]), float(following["y1"])),
                    "size": max(float(current["size"]), float(following["size"])),
                    "bold": bool(current["bold"] or following["bold"]),
                    "merged_heading": True,
                }
                index += 1
        merged.append(current)
        index += 1
    return merged


def extract_pdf(path: Path) -> ExtractedDocument:
    blocks: list[TextBlock] = []
    warnings: list[str] = []
    order = 0
    document = fitz.open(path)
    page_count = document.page_count
    try:
        all_font_sizes: list[float] = []
        raw_pages: list[list[dict]] = []
        for page_number, page in enumerate(document, start=1):
            page_lines: list[dict] = []
            payload = page.get_text("dict", sort=True)
            page_width = float(page.rect.width)
            page_height = float(page.rect.height)
            for raw_block in payload.get("blocks", []):
                if raw_block.get("type") != 0:
                    continue
                for line in raw_block.get("lines", []):
                    spans = line.get("spans", [])
                    text = _clean_text(
                        "".join(str(span.get("text", "")) for span in spans)
                    )
                    if not text:
                        continue
                    sizes = [
                        float(span.get("size", 0) or 0)
                        for span in spans
                        if span.get("text")
                    ]
                    size = max(sizes, default=0.0)
                    flags = [int(span.get("flags", 0) or 0) for span in spans]
                    bold = any(flag & 16 for flag in flags) or any(
                        "bold" in str(span.get("font", "")).casefold()
                        for span in spans
                    )
                    bbox = line.get("bbox") or raw_block.get("bbox") or (0, 0, 0, 0)
                    all_font_sizes.extend(size_value for size_value in sizes if size_value)
                    page_lines.append(
                        {
                            "text": text,
                            "page": page_number,
                            "size": size,
                            "bold": bold,
                            "x0": float(bbox[0]),
                            "y0": float(bbox[1]),
                            "x1": float(bbox[2]),
                            "y1": float(bbox[3]),
                            "page_width": page_width,
                            "page_height": page_height,
                        }
                    )
            raw_pages.append(_merge_split_headings(page_lines))

        body_size = median(all_font_sizes) if all_font_sizes else 0.0
        repeated_boundaries: dict[str, set[int]] = {}
        if page_count > 2:
            for page_number, page_lines in enumerate(raw_pages, start=1):
                for line in page_lines:
                    in_boundary = (
                        float(line["y0"]) <= float(line["page_height"]) * 0.09
                        or float(line["y1"]) >= float(line["page_height"]) * 0.91
                    )
                    key = _repeat_key(line["text"])
                    if in_boundary and len(key) >= 4 and len(line["text"]) < 160:
                        repeated_boundaries.setdefault(key, set()).add(page_number)

        repeat_threshold = max(3, round(page_count * 0.35))
        for page_lines in raw_pages:
            for line in page_lines:
                key = _repeat_key(line["text"])
                is_repeated_boundary = (
                    page_count > 2
                    and len(repeated_boundaries.get(key, set())) >= repeat_threshold
                )
                if is_repeated_boundary:
                    continue
                order += 1
                size = float(line["size"])
                is_heading = bool(line["bold"]) or bool(line.get("merged_heading"))
                if body_size and size >= body_size * 1.07:
                    is_heading = True
                blocks.append(
                    TextBlock(
                        text=line["text"],
                        page=int(line["page"]),
                        order=order,
                        style="Heading" if is_heading else "",
                        font_size=round(size, 2),
                        bold=bool(line["bold"]),
                        source_kind=(
                            "pdf_merged_heading"
                            if line.get("merged_heading")
                            else "pdf_line"
                        ),
                        x0=round(float(line["x0"]), 2),
                        y0=round(float(line["y0"]), 2),
                        x1=round(float(line["x1"]), 2),
                        y1=round(float(line["y1"]), 2),
                        page_width=round(float(line["page_width"]), 2),
                        page_height=round(float(line["page_height"]), 2),
                    )
                )
    finally:
        document.close()

    character_count = sum(len(block.text) for block in blocks)
    if character_count < max(250, page_count * 80):
        warnings.append(
            "Le document contient très peu de texte exploitable. Il peut s'agir "
            "d'un PDF scanné nécessitant un OCR."
        )
    return ExtractedDocument(path, blocks, page_count, character_count, warnings)


def _iter_docx_blocks(parent: DocumentObject) -> Iterable[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _is_toc_style(style_name: str) -> bool:
    # Word's automatic table-of-contents field applies built-in "TOC 1".."TOC 9"
    # paragraph styles (the underlying styleId stays this canonical English
    # name even in a French install). Each entry's plain text is the heading
    # title with its page number appended (the dot leader is a tab, not
    # literal dots), so left unfiltered these are indistinguishable from real
    # numbered headings and get extracted as duplicate, page-number-suffixed
    # phantom items ahead of the real content.
    normalized = style_name.strip().casefold()
    return bool(re.match(r"^toc\s*\d*$", normalized))


def extract_docx(path: Path) -> ExtractedDocument:
    document = Document(path)
    blocks: list[TextBlock] = []
    warnings: list[str] = []
    order = 0
    for element in _iter_docx_blocks(document):
        if isinstance(element, Paragraph):
            style_name = str(element.style.name if element.style else "")
            if _is_toc_style(style_name):
                continue
            text = _clean_text(element.text)
            if not text:
                continue
            order += 1
            runs = [run for run in element.runs if run.text.strip()]
            sizes = [
                float(run.font.size.pt)
                for run in runs
                if run.font.size is not None
            ]
            blocks.append(
                TextBlock(
                    text=text,
                    page=None,
                    order=order,
                    style=style_name,
                    font_size=max(sizes, default=0.0),
                    bold=bool(runs)
                    and sum(1 for run in runs if run.bold) >= len(runs) / 2,
                    source_kind="paragraph",
                )
            )
        else:
            for row in element.rows:
                cells = [_clean_text(cell.text) for cell in row.cells]
                text = " | ".join(cell for cell in cells if cell)
                if not text:
                    continue
                order += 1
                blocks.append(
                    TextBlock(
                        text=text,
                        page=None,
                        order=order,
                        style="Table",
                        source_kind="table_row",
                    )
                )

    character_count = sum(len(block.text) for block in blocks)
    if not blocks:
        warnings.append("Le document Word ne contient aucun paragraphe exploitable.")
    return ExtractedDocument(path, blocks, 0, character_count, warnings)


def extract_document(path: Path) -> ExtractedDocument:
    suffix = path.suffix.casefold()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    raise ValueError(f"Format non pris en charge : {path.suffix}")
